from datetime import datetime
from pathlib import Path

from docx import Document


def gerar_relatorio_word(parametros, artigos):
    document = Document()

    document.add_heading('Relatorio PRISMA Review Robot', level=1)

    document.add_paragraph(f"Data de execucao: {parametros.get('data_execucao', '')}")
    document.add_paragraph(f"Tema: {parametros.get('tema', '')}")
    document.add_paragraph(f"Ano inicial: {parametros.get('ano_inicial', '')}")
    document.add_paragraph(f"Ano final: {parametros.get('ano_final', '')}")
    document.add_paragraph(f"Tipo de revisao: {parametros.get('tipo_revisao', '')}")
    document.add_paragraph(f"Estrategia de busca: {parametros.get('query_geral', parametros.get('query_pubmed', ''))}")
    document.add_paragraph(f"Total de artigos processados: {len(artigos)}")

    document.add_heading('Artigos', level=2)

    if not artigos:
        document.add_paragraph('Nenhum artigo encontrado ou processado.')
    else:
        for idx, artigo in enumerate(artigos, start=1):
            titulo = artigo.get('Titulo', artigo.get('title', ''))
            autores = artigo.get('Autores', '')
            ano = artigo.get('Ano', '')
            revista = artigo.get('Revista', artigo.get('journal', ''))
            doi = artigo.get('DOI', '')
            link = artigo.get('Link', artigo.get('URL', ''))

            document.add_heading(f'Artigo {idx}', level=3)
            document.add_paragraph(f'Titulo: {titulo}')
            if autores:
                document.add_paragraph(f'Autores: {autores}')
            if ano:
                document.add_paragraph(f'Ano: {ano}')
            if revista:
                document.add_paragraph(f'Revista: {revista}')
            if doi:
                document.add_paragraph(f'DOI: {doi}')
            if link:
                document.add_paragraph(f'Link: {link}')

            resumo = artigo.get('Resumo', artigo.get('abstract', ''))
            if resumo:
                document.add_paragraph('Resumo:')
                document.add_paragraph(resumo)

            if idx >= 10:
                document.add_paragraph('...')
                document.add_paragraph('Apenas os 10 primeiros artigos foram incluidos no relatorio para manter o documento legivel.')
                break

    data_execucao = parametros.get('data_execucao', datetime.now().strftime('%Y-%m-%d_%H%M%S'))
    nome_arquivo = f"relatorio_prisma_robot_{data_execucao.replace(' ', '_').replace(':', '-')}.docx"
    caminho = Path('outputs') / nome_arquivo
    caminho.parent.mkdir(parents=True, exist_ok=True)

    document.save(caminho)

    print(f"Relatorio Word gerado em: {caminho}")
    return str(caminho)
