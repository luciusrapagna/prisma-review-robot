import shutil
import glob

try:
    from exportacao_prisma_elegante import exportar_master_prisma as exportar_master_prisma_elegante
except ImportError:
    from src.exportacao_prisma_elegante import exportar_master_prisma as exportar_master_prisma_elegante

from database.historico import (
    salvar_revisao,
    listar_revisoes,
    limpar_historico
)

from exports.abnt import gerar_referencias_abnt

from prisma_flow.fluxograma import gerar_fluxograma_prisma
from prisma.filtro_qualidade import filtrar_artigos_confiaveis, deduplicar_por_melhor_registro, auditar_descartes_openalex


import streamlit as st
from components.sidebar import render_sidebar_manutencao

st.set_page_config(
    page_title="PRISMA Review Robot",
    page_icon="📚",
    layout="wide"
)

st.title("📚 PRISMA Review Robot")
render_sidebar_manutencao()




st.subheader("ATHENA Scientific — Revisões Sistemáticas Inteligentes")

st.markdown("""
Sistema automatizado para:
- busca bibliográfica multibase
- remoção de duplicatas
- ranking semântico por IA
- geração automática de relatórios PRISMA
- exportação Word, tabelas e RIS Zotero
""")


historico = listar_revisoes()

with st.sidebar:
    st.header("📚 Histórico ATHENA")

    if historico:
        for item in historico[:10]:
            st.caption(
                f"{item[1]} | {item[8]}"
            )
    else:
        st.info("Nenhuma revisão registrada.")

    st.divider()

    if st.button(
        "🗑️ Limpar histórico",
        use_container_width=True,
        type="secondary"
    ):
        limpar_historico()
        st.success("Histórico apagado com sucesso.")
        st.rerun()

st.divider()


from datetime import datetime
import os
import hashlib
import pandas as pd

from buscadores.pubmed import executar_busca_pubmed
from buscadores.crossref import executar_busca_crossref
from buscadores.scielo import executar_busca_scielo
from buscadores.lilacs import executar_busca_lilacs
from buscadores.google_scholar_seguro import salvar_orientacoes_google_scholar
from buscadores.busca_academica_ampliada import executar_busca_academica_ampliada

from prisma.duplicates import remover_duplicatas

from ia.ranking_semantico import (
    calcular_similaridade,
    salvar_ranking_semantico
)

from engines.motor_booleano_master import gerar_estrategias

from outputs.word_writer import gerar_relatorio_word


def criar_pastas():
    pastas = [
        "data/raw",
        "data/processed",
        "outputs/tables",
        "outputs/figures",
        "outputs/references",
        "logs"
    ]

    for pasta in pastas:
        os.makedirs(pasta, exist_ok=True)


def coletar_parametros():
    import streamlit as st
    from datetime import date

    tema = st.text_input("Digite o tema da revisão", value="dengue")

    tipo_revisao = st.selectbox(
        "Tipo de revisão",
        ["Revisão sistemática", "Revisão de escopo", "Revisão integrativa", "Revisão narrativa"],
        index=1
    )

    estrategias = gerar_estrategias(
        tema=tema,
        tipo_revisao=tipo_revisao,
    )

    query_pubmed_automatica = estrategias["query_pubmed"]
    query_scielo_automatica = estrategias["query_scielo"]
    query_lilacs_automatica = estrategias["query_lilacs"]
    query_geral_automatica = estrategias["query_geral"]

    st.caption(
        "Estratégias automáticas geradas pelo Motor Booleano "
        "Master ATHENA."
    )

    aba_pubmed, aba_latam, aba_geral = st.tabs(
        ["PubMed", "SciELO/LILACS", "Busca geral"]
    )

    chave_base = hashlib.md5(
        f"{tema}|{tipo_revisao}".encode("utf-8")
    ).hexdigest()

    with aba_pubmed:
        query_pubmed = st.text_area(
            "Estratégia PubMed",
            value=query_pubmed_automatica,
            height=300,
            key=f"query_pubmed_{chave_base}",
        )

    with aba_latam:
        query_scielo = st.text_area(
            "Estratégia SciELO",
            value=query_scielo_automatica,
            height=300,
            key=f"query_scielo_{chave_base}",
        )

        query_lilacs = st.text_area(
            "Estratégia LILACS/BVS",
            value=query_lilacs_automatica,
            height=300,
            key=f"query_lilacs_{chave_base}",
        )

    with aba_geral:
        query_geral = st.text_area(
            "Estratégia de busca geral",
            value=query_geral_automatica,
            height=300,
            key=f"query_geral_{chave_base}",
        )

    relatorio_motor = estrategias.get("relatorio", {})
    conceitos_motor = relatorio_motor.get(
        "conceitos_identificados",
        [],
    )

    if conceitos_motor:
        st.info(
            "Conceitos identificados: "
            + ", ".join(conceitos_motor)
        )
    else:
        st.warning(
            "Nenhum conceito do vocabulário interno foi identificado. "
            "O motor aplicou uma estratégia conservadora baseada no tema."
        )

    col1, col2, col3 = st.columns(3)
    with col1:
        ano_inicial = st.number_input("Ano inicial", min_value=1900, max_value=2100, value=2014, step=1)
    with col2:
        ano_final = st.number_input("Ano final", min_value=1900, max_value=2100, value=2026, step=1)
    with col3:
        max_artigos = st.number_input("Número máximo de artigos", min_value=1, max_value=1000, value=30, step=1)

    data_execucao = st.date_input("Data de execução", value=date.today())

    bases = st.multiselect(
        "Bases de dados",
        ["PubMed", "Scopus", "Web of Science", "SciELO", "LILACS",
        "Google Acadêmico", "Busca Acadêmica Ampliada", "BVS"],
        default=["PubMed", "SciELO", "LILACS", "Google Acadêmico", "Busca Acadêmica Ampliada"]
    )

    idioma = st.multiselect(
        "Idiomas",
        ["Português", "Inglês", "Espanhol"],
        default=["Português", "Inglês", "Espanhol"]
    )

    tipo_estudo = st.text_input("Tipo de estudo", value=tipo_revisao)

    if not tema:
        st.warning("Informe o tema da revisão para continuar.")
        st.stop()

    return {
        "tema": tema,
        "tipo_revisao": tipo_revisao,
        "query_geral": query_geral,
        "query_pubmed": query_pubmed,
        "query_scielo": query_scielo,
        "query_lilacs": query_lilacs,
        "ano_inicial": int(ano_inicial),
        "ano_final": int(ano_final),
        "max_artigos": int(max_artigos),
        "bases": ", ".join(bases),
        "idioma": ", ".join(idioma),
        "tipo_estudo": tipo_estudo,
        "data_execucao": data_execucao.strftime("%d/%m/%Y")
    }


def perguntar_similaridade():
    import streamlit as st

    valor = st.number_input(
        "Digite a similaridade mínima desejada",
        min_value=0.0,
        max_value=1.0,
        value=0.35,
        step=0.05
    )

    return float(valor)


def gerar_tabela_parametros(parametros, similaridade_minima):
    dados = {
        "Campo": [
            "Tema",
            "Ano inicial",
            "Ano final",
            "EstratAgia de busca",
            "MAximo de artigos por base",
            "Tipo de revisAo",
            "Similaridade mAnima",
            "Data de execuAAo"
        ],
        "InformaAAo": [
            parametros["tema"],
            parametros["ano_inicial"],
            parametros["ano_final"],
            parametros["query_geral"],
            parametros["max_artigos"],
            parametros["tipo_revisao"],
            similaridade_minima,
            parametros["data_execucao"]
        ]
    }

    df = pd.DataFrame(dados)

    caminho = "outputs/tables/parametros_revisao.xlsx"
    df.to_excel(caminho, index=False)

    print(f"\nTabela de parAmetros gerada em: {caminho}")


def salvar_tabela_consolidada(artigos):
    caminho = "outputs/tables/tabela_consolidada_multibase.xlsx"

    if not artigos:
        df = pd.DataFrame(columns=[
            "Base",
            "PMID",
            "TAtulo",
            "Autores",
            "Ano",
            "Revista",
            "DOI",
            "Resumo",
            "Link"
        ])
    else:
        df = pd.DataFrame(artigos)

    df.to_excel(caminho, index=False)

    print(f"\nTabela consolidada multibase salva em: {caminho}")


def gerar_descricao_figura_prisma(
    parametros,
    total_pubmed,
    total_crossref,
    total_scielo,
    total_lilacs,
    total_identificados,
    total_sem_duplicatas,
    total_apos_similaridade,
    similaridade_minima
):
    duplicatas = total_identificados - total_sem_duplicatas
    excluidos_por_similaridade = total_sem_duplicatas - total_apos_similaridade

    texto = f"""
DescriAAo sugerida para figura PRISMA:

Fluxograma representando o processo de identificaAAo, triagem, elegibilidade e inclusAo dos estudos da revisAo intitulada "{parametros['tema']}".

A busca bibliogrAfica foi realizada nas bases PubMed, Crossref, SciELO e LILACS/BVS, considerando publicaAAes entre {parametros['ano_inicial']} e {parametros['ano_final']}.

A estratAgia de busca utilizada foi:

{parametros['query_geral']}

Na etapa de identificaAAo, foram recuperados:
- PubMed: {total_pubmed} registros;
- Crossref: {total_crossref} registros;
- SciELO: {total_scielo} registros;
- LILACS/BVS: {total_lilacs} registros.

O total inicial identificado foi de {total_identificados} registros.

ApAs a remoAAo automAtica de duplicatas, permaneceram {total_sem_duplicatas} registros Anicos. Foram removidos {duplicatas} registros duplicados.

ApAs a triagem semAntica automatizada, utilizando similaridade mAnima de {similaridade_minima}, permaneceram {total_apos_similaridade} registros. Foram excluAdos {excluidos_por_similaridade} registros por baixa similaridade com o tema da revisAo.

A figura PRISMA deverA apresentar:
1. registros identificados em cada base;
2. total de registros identificados;
3. registros removidos por duplicidade;
4. registros triados por similaridade semAntica;
5. registros excluAdos por baixa similaridade;
6. registros mantidos para leitura de tAtulo e resumo;
7. textos completos avaliados para elegibilidade;
8. textos completos excluAdos com justificativa;
9. estudos finais incluAdos na sAntese qualitativa e/ou quantitativa.

SugestAo visual:
Construir um fluxograma vertical em quatro etapas principais: IdentificaAAo, Triagem, Elegibilidade e InclusAo. Cada base de dados pode aparecer em uma caixa lateral na etapa de identificaAAo, convergindo para o total de registros identificados. Em seguida, inserir a caixa de remoAAo de duplicatas, seguida da triagem semAntica, avaliaAAo de elegibilidade e inclusAo final.

Essa descriAAo pode ser utilizada como base para criaAAo da figura no PowerPoint, Canva, CorelDRAW, BioRender ou outro software grAfico.
"""

    caminho = "outputs/figures/descricao_figura_prisma.txt"

    with open(caminho, "w", encoding="utf-8") as arquivo:
        arquivo.write(texto)

    print(f"\nDescriAAo da figura PRISMA gerada em: {caminho}")


def gerar_ris_zotero(artigos):
    caminho = "outputs/references/referencias_multibase.ris"

    with open(caminho, "w", encoding="utf-8") as arquivo:
        for artigo in artigos:
            arquivo.write("TY  - JOUR\n")
            arquivo.write(f"TI  - {artigo.get('TAtulo', '')}\n")
            arquivo.write(f"PY  - {artigo.get('Ano', '')}\n")
            arquivo.write(f"JO  - {artigo.get('Revista', '')}\n")
            arquivo.write(f"DO  - {artigo.get('DOI', '')}\n")
            arquivo.write(f"UR  - {artigo.get('Link', '')}\n")

            autores = artigo.get("Autores", "")
            if autores:
                if isinstance(autores, list):
                    for autor in autores:
                        arquivo.write(f"AU  - {autor}\n")
                else:
                    for autor in str(autores).split("; "):
                        if autor.strip():
                            arquivo.write(f"AU  - {autor.strip()}\n")

            arquivo.write("ER  -\n\n")

    print(f"\nArquivo RIS para Zotero gerado em: {caminho}")


def executar_buscas(parametros):
    artigos_pubmed = executar_busca_pubmed(
        query=parametros["query_pubmed"],
        ano_inicial=parametros["ano_inicial"],
        ano_final=parametros["ano_final"],
        max_artigos=parametros["max_artigos"],
        tema=parametros["tema"]
    )

    artigos_crossref = executar_busca_crossref(
        query=parametros["query_geral"],
        ano_inicial=parametros["ano_inicial"],
        ano_final=parametros["ano_final"],
        max_artigos=parametros["max_artigos"]
    )

    artigos_scielo = executar_busca_scielo(
        query=parametros["query_geral"],
        ano_inicial=parametros["ano_inicial"],
        ano_final=parametros["ano_final"],
        max_artigos=parametros["max_artigos"]
    )

    artigos_lilacs = executar_busca_lilacs(
        query=parametros["query_geral"],
        ano_inicial=parametros["ano_inicial"],
        ano_final=parametros["ano_final"],
        max_artigos=parametros["max_artigos"]
    )

    artigos_ampliada = []

    if "Busca Acadêmica Ampliada" in parametros.get("bases", []):
        artigos_ampliada = executar_busca_academica_ampliada(
            query=parametros.get("query_geral", parametros["query_pubmed"]),
            query_ingles=parametros.get("query_pubmed", parametros.get("query_geral")),
            ano_inicial=parametros["ano_inicial"],
            ano_final=parametros["ano_final"],
            max_artigos=parametros["max_artigos"]
        )

    return artigos_pubmed, artigos_crossref, artigos_scielo, artigos_lilacs, artigos_ampliada




def render_google_academico_seguro(parametros):
    bases = parametros.get("bases", [])

    if "Google Acadêmico" not in bases:
        return

    query_google = (
        parametros.get("query_geral")
        or parametros.get("query_latam")
        or parametros.get("query_pubmed")
        or parametros.get("tema")
    )

    caminho_google, link_google = salvar_orientacoes_google_scholar(
        query=query_google,
        ano_inicial=parametros.get("ano_inicial"),
        ano_final=parametros.get("ano_final")
    )

    st.subheader("🔎 Google Acadêmico — busca segura")
    st.info(
        "O Google Acadêmico não possui API pública oficial estável. "
        "Por isso, o ATHENA gera um link de busca seguro para abrir no navegador. "
        "Depois, os resultados podem ser importados para auditoria, deduplicação e ranking."
    )

    st.link_button("Abrir busca no Google Acadêmico", link_google)

    with open(caminho_google, "r", encoding="utf-8") as f:
        conteudo = f.read()

    st.download_button(
        "Baixar estratégia do Google Acadêmico",
        data=conteudo,
        file_name="busca_google_academico_athena.txt",
        mime="text/plain"
    )

    st.code(query_google, language="text")



def main():
    criar_pastas()

    parametros = coletar_parametros()

    similaridade_minima = perguntar_similaridade()

    import streamlit as st

    # Google Acadêmico — bloco seguro visível antes da execução
    render_google_academico_seguro(parametros)

    executar = st.button("🚀 Executar revisão")

    if not executar:
        st.stop()

    gerar_tabela_parametros(
        parametros,
        similaridade_minima
    )

    (
        artigos_pubmed,
        artigos_crossref,
        artigos_scielo,
        artigos_lilacs,
        artigos_ampliada
    ) = executar_buscas(parametros)

    todos_artigos = (
        artigos_pubmed
        + artigos_crossref
        + artigos_scielo
        + artigos_lilacs
        + artigos_ampliada
    )

    print("\n" + "=" * 70)
    print("ATHENA PRISMA — DIAGNÓSTICO DAS BASES")
    print("=" * 70)
    print(f"Bases selecionadas na interface: {parametros.get('bases', [])}")
    print(f"PubMed: {len(artigos_pubmed)}")
    print(f"Crossref: {len(artigos_crossref)}")
    print(f"SciELO: {len(artigos_scielo)}")
    print(f"LILACS: {len(artigos_lilacs)}")
    print(f"Busca Acadêmica Ampliada: {len(artigos_ampliada)}")
    print(f"TOTAL BRUTO: {len(todos_artigos)}")

    from collections import Counter
    distribuicao_bases = Counter(
        str(a.get("Base") or a.get("base") or "Não informado")
        for a in todos_artigos
        if isinstance(a, dict)
    )
    print(f"Distribuição real: {dict(distribuicao_bases)}")
    print("=" * 70 + "\n")

    artigos_confiaveis, artigos_descartados_qualidade = filtrar_artigos_confiaveis(
        todos_artigos
    )

    try:
        caminho_auditoria_openalex = auditar_descartes_openalex(
            artigos_originais=todos_artigos,
            artigos_confiaveis=artigos_confiaveis,
            artigos_descartados=artigos_descartados_qualidade,
            caminho="outputs/tables/auditoria_openalex_descartes.xlsx",
        )
        print(
            f"Auditoria OpenAlex gerada em: "
            f"{caminho_auditoria_openalex}"
        )
    except Exception as erro_openalex:
        import traceback
        print(
            f"ERRO AO GERAR AUDITORIA OPENALEX: "
            f"{erro_openalex}"
        )
        traceback.print_exc()

    artigos_confiaveis = deduplicar_por_melhor_registro(
        artigos_confiaveis
    )

    salvar_tabela_consolidada(artigos_confiaveis)

    total_pubmed = len(artigos_pubmed)
    total_crossref = len(artigos_crossref)
    total_scielo = len(artigos_scielo)
    total_lilacs = len(artigos_lilacs)
    total_ampliada = len(artigos_ampliada)
    total_identificados = len(todos_artigos)

    artigos_sem_duplicatas = remover_duplicatas(artigos_confiaveis)

    total_sem_duplicatas = len(artigos_sem_duplicatas)

    artigos_rankeados = calcular_similaridade(
        tema=parametros["tema"],
        artigos=artigos_sem_duplicatas,
        similaridade_minima=similaridade_minima
    )

    total_apos_similaridade = len(artigos_rankeados)

    salvar_ranking_semantico(artigos_rankeados)

    caminho_relatorio = gerar_relatorio_word(
        parametros,
        artigos_rankeados
    )

    gerar_descricao_figura_prisma(
        parametros=parametros,
        total_pubmed=total_pubmed,
        total_crossref=total_crossref,
        total_scielo=total_scielo,
        total_lilacs=total_lilacs,
        total_identificados=total_identificados,
        total_sem_duplicatas=total_sem_duplicatas,
        total_apos_similaridade=total_apos_similaridade,
        similaridade_minima=similaridade_minima
    )

    gerar_ris_zotero(artigos_rankeados)

    gerar_referencias_abnt(artigos_rankeados)

    svg_fluxo, pdf_fluxo = gerar_fluxograma_prisma(
        total_identificados=total_identificados,
        total_sem_duplicatas=total_sem_duplicatas,
        total_apos_similaridade=total_apos_similaridade
    )

    arquivos_master = exportar_master_prisma_elegante(
        artigos_totais=artigos_confiaveis,
        artigos_incluidos=artigos_rankeados,
        svg_fluxo=svg_fluxo,
        jpg_fluxo=None,
        pdf_fluxo=pdf_fluxo,
    )

    salvar_revisao(
        parametros=parametros,
        total_identificados=total_identificados,
        total_pos_duplicatas=total_sem_duplicatas,
        total_pos_similaridade=total_apos_similaridade
    )


    import streamlit as st
    from pathlib import Path
    import zipfile
    from io import BytesIO

    st.success("✅ Robô PRISMA executado com sucesso!")

    st.write(f"**PubMed:** {total_pubmed} registros")
    st.write(f"**Crossref:** {total_crossref} registros")
    st.write(f"**SciELO:** {total_scielo} registros")
    st.write(f"**LILACS/BVS:** {total_lilacs} registros")
    st.write(f"**Total identificado:** {total_identificados} registros")
    st.write(f"**Total após duplicatas:** {total_sem_duplicatas} registros")
    st.write(f"**Total após similaridade:** {total_apos_similaridade} registros")
    st.write(f"**Similaridade mínima usada:** {similaridade_minima}")

    st.subheader("📊 Fluxograma PRISMA")

    st.image(svg_fluxo)

    with open(svg_fluxo, "rb") as f:
        st.download_button(
            "🖼️ Baixar fluxograma SVG",
            f,
            file_name="fluxograma_prisma.svg",
            mime="image/svg+xml"
        )

    with open(pdf_fluxo, "rb") as f:
        st.download_button(
            "📄 Baixar fluxograma PDF",
            f,
            file_name="fluxograma_prisma.pdf",
            mime="application/pdf"
        )

    st.subheader("📁 Arquivos gerados")

    caminho_relatorio = Path(caminho_relatorio)

    if caminho_relatorio.exists():
        with open(caminho_relatorio, "rb") as f:
            st.download_button(
                "📄 Baixar relatório Word",
                f,
                file_name=caminho_relatorio.name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    memoria_zip = BytesIO()

    with zipfile.ZipFile(memoria_zip, "w", zipfile.ZIP_DEFLATED) as zipf:
        for pasta in ["outputs/tables", "outputs/figures", "outputs/references"]:
            p = Path(pasta)
            if p.exists():
                for arquivo_saida in p.rglob("*"):
                    if arquivo_saida.is_file():
                        zipf.write(arquivo_saida, arquivo_saida.as_posix())

    memoria_zip.seek(0)

    st.download_button(
        "📦 Baixar todos os arquivos gerados",
        memoria_zip,
        file_name="prisma_outputs.zip",
        mime="application/zip"
    )


    st.subheader("📦 Exportação master ATHENA PRISMA")

    for nome, caminho in arquivos_master.items():
        caminho = Path(caminho)
        if caminho.exists():
            with open(caminho, "rb") as f:
                st.download_button(
                    label=f"⬇️ Baixar {nome}",
                    data=f,
                    file_name=caminho.name,
                    mime="application/octet-stream"
                )






if __name__ == "__main__":
    main()


from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _autor_curto_prisma(autores):
    if not autores:
        return ""

    if isinstance(autores, str):
        lista = [a.strip() for a in autores.replace(";", ",").split(",") if a.strip()]
    else:
        lista = autores

    sobrenomes = []
    for autor in lista:
        partes = str(autor).strip().split()
        if partes:
            sobrenomes.append(partes[-1])

    if len(sobrenomes) == 1:
        return sobrenomes[0]
    if len(sobrenomes) == 2:
        return f"{sobrenomes[0]} & {sobrenomes[1]}"
    return f"{sobrenomes[0]} et al."


def _campo_prisma(artigo, *nomes, padrao=""):
    if not isinstance(artigo, dict):
        return padrao
    for nome in nomes:
        valor = artigo.get(nome)
        if valor not in [None, ""]:
            return valor
    return padrao


def exportar_excel_master_elegante(artigos, pasta_saida="outputs", artigos_incluidos=None):
    Path(pasta_saida).mkdir(parents=True, exist_ok=True)

    caminho = Path(pasta_saida) / "ATHENA_PRISMA_Master.xlsx"

    wb = Workbook()

    def preencher_aba(ws, lista_artigos, status_padrao):
        ws.title = ws.title[:31]

        colunas = [
            "ID", "Autor", "Ano", "Título", "Revista",
            "DOI", "Link", "Base", "Status", "Similaridade"
        ]
        ws.append(colunas)

        for i, artigo in enumerate(lista_artigos or [], start=1):
            autores = _campo_prisma(artigo, "Autores", "autores", "authors")

            ws.append([
                i,
                _autor_curto_prisma(autores),
                _campo_prisma(artigo, "Ano", "ano", "year"),
                _campo_prisma(artigo, "Titulo", "Título", "titulo", "title"),
                _campo_prisma(artigo, "Revista", "revista", "journal", "fonte"),
                _campo_prisma(artigo, "DOI", "doi"),
                _campo_prisma(artigo, "Link", "link", "url", "pubmed_url"),
                _campo_prisma(artigo, "Base", "base", "database", padrao="Sem base"),
                _campo_prisma(artigo, "status", "Status", padrao=status_padrao),
                _campo_prisma(
                    artigo,
                    "similaridade",
                    "Similaridade",
                    "score_similaridade",
                    padrao=""
                ),
            ])

        header_fill = PatternFill("solid", fgColor="1F4E78")
        header_font = Font(color="FFFFFF", bold=True)
        thin = Side(border_style="thin", color="D9E2F3")

        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)

        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(vertical="center", wrap_text=True)
                cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)

        for row in range(2, ws.max_row + 1):
            doi_cell = ws[f"F{row}"]
            link_cell = ws[f"G{row}"]

            if doi_cell.value:
                doi_cell.hyperlink = f"https://doi.org/{doi_cell.value}"
                doi_cell.style = "Hyperlink"

            if link_cell.value:
                link_cell.hyperlink = str(link_cell.value)
                link_cell.style = "Hyperlink"

        ws.freeze_panes = "A2"

        larguras = {
            "A": 8, "B": 22, "C": 10, "D": 70, "E": 35,
            "F": 28, "G": 45, "H": 22, "I": 18, "J": 14
        }

        for col, largura in larguras.items():
            ws.column_dimensions[col].width = largura

    ws1 = wb.active
    ws1.title = "Identificados"
    preencher_aba(ws1, artigos, "Identificado")

    ws2 = wb.create_sheet("Incluídos após triagem")
    preencher_aba(ws2, artigos_incluidos if artigos_incluidos is not None else artigos, "Incluído")

    wb.save(caminho)

    return caminho

def exportar_word_tabela_artigo(artigos, pasta_saida="outputs"):
    Path(pasta_saida).mkdir(parents=True, exist_ok=True)

    caminho = Path(pasta_saida) / "ATHENA_PRISMA_Tabela_Artigo.docx"

    doc = Document()

    titulo = doc.add_heading("ATHENA PRISMA REVIEW ROBOT", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = doc.add_heading("Tabela de Estudos Incluídos", level=2)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(f"Estudos incluídos (n = {len(artigos)}).")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = "Table Grid"

    headers = ["Autor", "Ano", "Título", "Revista"]
    for i, h in enumerate(headers):
        tabela.rows[0].cells[i].text = h

    for cell in tabela.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

    for artigo in artigos:
        autores = _campo_prisma(artigo, "autores", "authors")
        row = tabela.add_row().cells
        row[0].text = _autor_curto_prisma(autores)
        row[1].text = str(_campo_prisma(artigo, "ano", "year"))
        row[2].text = str(_campo_prisma(artigo, "titulo", "title"))
        row[3].text = str(_campo_prisma(artigo, "revista", "journal", "fonte"))

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

    doc.save(caminho)
    return str(caminho)


def exportar_master_prisma(artigos_totais, artigos_incluidos, svg_fluxo=None, jpg_fluxo=None, pdf_fluxo=None):
    arquivos = {}

    arquivos["excel_master"] = exportar_excel_master_elegante(
        artigos_totais,
        artigos_incluidos=artigos_incluidos
    )
    arquivos["word_tabela_artigo"] = exportar_word_tabela_artigo(artigos_incluidos)

    if svg_fluxo:
        arquivos["fluxograma_svg"] = svg_fluxo

    if pdf_fluxo:
        arquivos["fluxograma_pdf"] = pdf_fluxo

    return arquivos
