from collections import Counter
import re
from pathlib import Path
from datetime import datetime

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def campo(artigo, *chaves, padrao=""):
    if not isinstance(artigo, dict):
        return padrao
    for chave in chaves:
        valor = artigo.get(chave)
        if valor not in (None, "", [], {}):
            return valor
    return padrao


def autor_abnt(autores):
    if not autores:
        return ""

    if isinstance(autores, str):
        lista = [a.strip() for a in autores.replace("|", ";").replace(",", ";").split(";") if a.strip()]
    else:
        lista = [str(a).strip() for a in autores if str(a).strip()]

    autores_formatados = []

    for autor in lista:
        partes = autor.strip().replace(".", "").split()
        if not partes:
            continue

        # Formato PubMed comum: "Mahajan R" ou "Davis DMR"
        sobrenome = partes[0].upper()
        iniciais = partes[1:]

        iniciais_formatadas = []
        for bloco in iniciais:
            for letra in bloco:
                if letra.isalpha():
                    iniciais_formatadas.append(f"{letra.upper()}.")

        if iniciais_formatadas:
            autores_formatados.append(f"{sobrenome}, {' '.join(iniciais_formatadas)}")
        else:
            autores_formatados.append(sobrenome)

    if len(autores_formatados) > 3:
        return f"{autores_formatados[0]} et al."

    return "; ".join(autores_formatados)


def normalizar_artigo(artigo):
    return {
        "autor": autor_abnt(campo(artigo, "autores", "Autores", "authors", "Authors", "author")),
        "ano": campo(artigo, "ano", "Ano", "year", "Year"),
        "titulo": campo(artigo, "titulo", "Titulo", "title", "Title", "Título", "titulo_artigo", "Título do artigo"),
        "revista": campo(artigo, "revista", "Revista", "journal", "Journal", "periódico", "periodico", "Periódico", "fonte"),
        "doi": campo(artigo, "doi", "DOI"),
        "link": campo(artigo, "link", "url", "URL", "pubmed_url"),
        "base": campo(artigo, "base", "database", "Base", padrao="PubMed"),
        "status": campo(artigo, "status", "Status", padrao="Incluído"),
    }


def exportar_excel_master(artigos, pasta_saida="outputs"):
    Path(pasta_saida).mkdir(parents=True, exist_ok=True)
    caminho = Path(pasta_saida) / "ATHENA_PRISMA_Master.xlsx"

    wb = Workbook()
    ws = wb.active
    ws.title = "Estudos incluídos"

    headers = ["ID", "Autor", "Ano", "Título", "Revista", "DOI", "Link", "Base", "Status"]
    ws.append(headers)

    for i, artigo in enumerate(artigos, 1):
        a = normalizar_artigo(artigo)
        ws.append([
            i,
            a["autor"],
            a["ano"],
            a["titulo"],
            a["revista"],
            a["doi"],
            a["link"],
            a["base"],
            a["status"],
        ])

    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    thin = Side(border_style="thin", color="D9E2F3")

    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(top=thin, bottom=thin, left=thin, right=thin)

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            cell.border = Border(top=thin, bottom=thin, left=thin, right=thin)

    for row in range(2, ws.max_row + 1):
        doi = ws[f"F{row}"].value
        url = ws[f"G{row}"].value

        if doi:
            ws[f"F{row}"].hyperlink = f"https://doi.org/{doi}"
            ws[f"F{row}"].style = "Hyperlink"

        if url:
            ws[f"G{row}"].hyperlink = str(url)
            ws[f"G{row}"].style = "Hyperlink"

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions

    widths = {
        "A": 8,
        "B": 18,
        "C": 10,
        "D": 55,
        "E": 32,
        "F": 30,
        "G": 42,
        "H": 14,
        "I": 14,
    }

    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    wb.save(caminho)
    return str(caminho)


def exportar_word_tabela_artigo(artigos, pasta_saida="outputs"):
    Path(pasta_saida).mkdir(parents=True, exist_ok=True)
    caminho = Path(pasta_saida) / "ATHENA_PRISMA_Tabela_Artigo.docx"

    doc = Document()

    titulo = doc.add_heading("ATHENA PRISMA REVIEW ROBOT", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = doc.add_heading("Tabela de Estudos Incluídos", level=2)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(f"Estudos incluídos (n = {len(artigos)}).")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = "Table Grid"

    headers = ["Autor", "Ano", "Título", "Revista"]
    for i, h in enumerate(headers):
        tabela.rows[0].cells[i].text = h

    for artigo in artigos:
        a = normalizar_artigo(artigo)
        row = tabela.add_row().cells
        row[0].text = str(a["autor"])
        row[1].text = str(a["ano"])
        row[2].text = str(a["titulo"])
        row[3].text = str(a["revista"])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

    doc.save(caminho)
    return str(caminho)


def exportar_master_prisma(
    artigos_totais,
    artigos_incluidos,
    svg_fluxo=None,
    jpg_fluxo=None,
    pdf_fluxo=None
):
    """
    Exportador oficial ATHENA PRISMA.

    Excel Master:
    - Aba 1: Estudos incluídos
    - Aba 2: Auditoria da busca
    """

    arquivos = {}

    print("\n===== ATHENA PRISMA — EXPORTAÇÃO MASTER =====")
    print(f"Total identificado recebido: {len(artigos_totais or [])}")
    print(f"Total incluído recebido: {len(artigos_incluidos or [])}")

    arquivos["excel_master"] = exportar_excel_master_elegante(
        artigos_totais,
        artigos_incluidos=artigos_incluidos
    )

    try:
        arquivos["word_tabela_artigo"] = exportar_word_tabela_artigo(
            artigos_incluidos
        )
    except Exception as e:
        print(f"AVISO Word: {e}")

    if svg_fluxo:
        arquivos["fluxograma_svg"] = svg_fluxo

    if jpg_fluxo:
        arquivos["fluxograma_jpg"] = jpg_fluxo

    if pdf_fluxo:
        arquivos["fluxograma_pdf"] = pdf_fluxo

    return arquivos


def _athena_get(artigo, *nomes, padrao=""):
    if not isinstance(artigo, dict):
        return padrao
    for nome in nomes:
        if nome in artigo and artigo.get(nome) not in [None, ""]:
            return artigo.get(nome)
    return padrao


def _athena_norm_texto(x):
    x = str(x or "").strip().lower()
    x = re.sub(r"\s+", " ", x)
    x = re.sub(r"[^\w\s]", "", x)
    return x.strip()


def _athena_norm_doi(x):
    x = str(x or "").strip()
    x = x.replace("https://doi.org/", "").replace("http://doi.org/", "")
    x = x.replace("doi:", "").replace("DOI:", "").strip()
    return x.lower()


def _athena_autor_curto(autores):
    if isinstance(autores, list):
        nomes = []
        for a in autores:
            if isinstance(a, dict):
                nome = (
                    a.get("name")
                    or a.get("display_name")
                    or f"{a.get('family', '')} {a.get('given', '')}".strip()
                )
                if nome:
                    nomes.append(str(nome).strip())
            elif str(a).strip():
                nomes.append(str(a).strip())
    else:
        bruto = str(autores or "").strip().replace("|", ";")
        nomes = [a.strip() for a in bruto.split(";") if a.strip()]

    if not nomes:
        return "Autores não informados"

    if len(nomes) <= 3:
        return "; ".join(nomes)

    return f"{nomes[0]} et al."


def _athena_linha_artigo(i, artigo, status_padrao):
    autores = _athena_get(artigo, "Autores", "autores", "authors")
    return [
        i,
        _athena_autor_curto(autores),
        _athena_get(artigo, "Ano", "ano", "year"),
        _athena_get(artigo, "Título", "Titulo", "titulo", "title"),
        _athena_get(artigo, "Revista", "revista", "journal", "fonte"),
        _athena_get(artigo, "DOI", "doi"),
        _athena_get(artigo, "Link", "link", "url", "pubmed_url"),
        _athena_get(artigo, "Base", "base", "database", padrao="Não informado"),
        _athena_get(artigo, "Status", "status", padrao=status_padrao),
        _athena_get(artigo, "Similaridade", "similaridade", "Score_Semantico", "Score_Ajustado_Multibase", "score", padrao=""),
    ]


def _athena_auditar_artigos(artigos):
    dois = [_athena_norm_doi(_athena_get(a, "DOI", "doi")) for a in artigos or []]
    titulos = [_athena_norm_texto(_athena_get(a, "Título", "Titulo", "titulo", "title")) for a in artigos or []]

    c_doi = Counter([d for d in dois if d])
    c_tit = Counter([t for t in titulos if t])

    linhas = []
    for i, artigo in enumerate(artigos or [], start=1):
        titulo = _athena_get(artigo, "Título", "Titulo", "titulo", "title")
        doi = _athena_get(artigo, "DOI", "doi")
        link = _athena_get(artigo, "Link", "link", "url", "pubmed_url")
        base = _athena_get(artigo, "Base", "base", "database", padrao="Não informado")

        nd = _athena_norm_doi(doi)
        nt = _athena_norm_texto(titulo)

        alertas = []

        if not str(titulo or "").strip():
            alertas.append("título ausente")
        if nt in ["title page", "untitled", "sem titulo", "sem título"] or len(nt) < 12:
            alertas.append("título genérico/suspeito")
        if not nd:
            alertas.append("DOI ausente")
        elif not re.match(r"^10\.\d{4,9}/\S+$", nd):
            alertas.append("DOI possivelmente inválido")
        if not str(link or "").strip():
            alertas.append("link ausente")
        if nd and c_doi[nd] > 1:
            alertas.append("duplicata por DOI")
        if nt and c_tit[nt] > 1:
            alertas.append("duplicata por título")
        if base in ["", "Não informado", None]:
            alertas.append("base não informada")

        linhas.append([
            i,
            base,
            titulo,
            doi,
            link,
            "Revisar" if alertas else "OK",
            "; ".join(alertas) if alertas else "Sem alertas",
        ])

    return linhas


def _athena_formatar_ws(ws):
    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    thin = Side(border_style="thin", color="D9E2F3")

    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions

    for col in range(1, ws.max_column + 1):
        letra = get_column_letter(col)
        if letra in ["D", "E", "G"]:
            ws.column_dimensions[letra].width = 45
        else:
            ws.column_dimensions[letra].width = 18


def exportar_excel_master_elegante(artigos, pasta_saida="outputs", artigos_incluidos=None):
    Path(pasta_saida).mkdir(parents=True, exist_ok=True)
    caminho = Path(pasta_saida) / "ATHENA_PRISMA_Master.xlsx"

    artigos_totais = artigos or []
    artigos_finais = artigos_incluidos if artigos_incluidos is not None else artigos_totais

    wb = Workbook()

    ws = wb.active
    ws.title = "Estudos incluídos"
    ws.append(["ID", "Autor", "Ano", "Título", "Revista", "DOI", "Link", "Base", "Status", "Similaridade"])

    for i, artigo in enumerate(artigos_finais or [], start=1):
        ws.append(_athena_linha_artigo(i, artigo, "Incluído"))

    _athena_formatar_ws(ws)

    for row in range(2, ws.max_row + 1):
        doi = ws[f"F{row}"].value
        link = ws[f"G{row}"].value
        if doi:
            ws[f"F{row}"].hyperlink = f"https://doi.org/{_athena_norm_doi(doi)}"
            ws[f"F{row}"].style = "Hyperlink"
        if link:
            ws[f"G{row}"].hyperlink = str(link)
            ws[f"G{row}"].style = "Hyperlink"

    ws2 = wb.create_sheet("Auditoria da busca")
    ws2.append(["ID", "Base", "Título", "DOI", "Link", "Status auditoria", "Alertas"])
    for linha in _athena_auditar_artigos(artigos_totais):
        ws2.append(linha)

    _athena_formatar_ws(ws2)

    wb.save(caminho)
    return str(caminho)


def exportar_master_prisma_elegante(artigos_totais, artigos_incluidos, svg_fluxo=None, jpg_fluxo=None, pdf_fluxo=None):
    arquivos = {}
    arquivos["excel_master"] = exportar_excel_master_elegante(
        artigos_totais,
        artigos_incluidos=artigos_incluidos
    )

    try:
        arquivos["word_tabela_artigo"] = exportar_word_tabela_artigo(artigos_incluidos)
    except Exception as e:
        print(f"Aviso: Word não exportado: {e}")

    if svg_fluxo:
        arquivos["fluxograma_svg"] = svg_fluxo
    if jpg_fluxo:
        arquivos["fluxograma_jpg"] = jpg_fluxo
    if pdf_fluxo:
        arquivos["fluxograma_pdf"] = pdf_fluxo

    return arquivos
