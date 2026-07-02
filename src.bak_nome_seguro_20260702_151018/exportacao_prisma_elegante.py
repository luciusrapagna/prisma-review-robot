from pathlib import Path
from datetime import datetime

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def campo(artigo, *chaves, padrao=""):
    if not isinstance(artigo, dict):
        return padrao
    for chave in chaves:
        valor = artigo.get(chave)
        if valor not in (None, "", [], {}):
            return valor
    return padrao


def autor_abnt(autores):
    if not autores:
        return ""

    if isinstance(autores, str):
        lista = [a.strip() for a in autores.replace("|", ";").replace(",", ";").split(";") if a.strip()]
    else:
        lista = [str(a).strip() for a in autores if str(a).strip()]

    autores_formatados = []

    for autor in lista:
        partes = autor.strip().replace(".", "").split()
        if not partes:
            continue

        # Formato PubMed comum: "Mahajan R" ou "Davis DMR"
        sobrenome = partes[0].upper()
        iniciais = partes[1:]

        iniciais_formatadas = []
        for bloco in iniciais:
            for letra in bloco:
                if letra.isalpha():
                    iniciais_formatadas.append(f"{letra.upper()}.")

        if iniciais_formatadas:
            autores_formatados.append(f"{sobrenome}, {' '.join(iniciais_formatadas)}")
        else:
            autores_formatados.append(sobrenome)

    if len(autores_formatados) > 3:
        return f"{autores_formatados[0]} et al."

    return "; ".join(autores_formatados)


def normalizar_artigo(artigo):
    return {
        "autor": autor_abnt(campo(artigo, "autores", "Autores", "authors", "Authors", "author")),
        "ano": campo(artigo, "ano", "Ano", "year", "Year"),
        "titulo": campo(artigo, "titulo", "Titulo", "title", "Title", "Título", "titulo_artigo", "Título do artigo"),
        "revista": campo(artigo, "revista", "Revista", "journal", "Journal", "periódico", "periodico", "Periódico", "fonte"),
        "doi": campo(artigo, "doi", "DOI"),
        "link": campo(artigo, "link", "url", "URL", "pubmed_url"),
        "base": campo(artigo, "base", "database", "Base", padrao="PubMed"),
        "status": campo(artigo, "status", "Status", padrao="Incluído"),
    }


def exportar_excel_master(artigos, pasta_saida="outputs"):
    Path(pasta_saida).mkdir(parents=True, exist_ok=True)
    caminho = Path(pasta_saida) / "ATHENA_PRISMA_Master.xlsx"

    wb = Workbook()
    ws = wb.active
    ws.title = "Estudos incluídos"

    headers = ["ID", "Autor", "Ano", "Título", "Revista", "DOI", "Link", "Base", "Status"]
    ws.append(headers)

    for i, artigo in enumerate(artigos, 1):
        a = normalizar_artigo(artigo)
        ws.append([
            i,
            a["autor"],
            a["ano"],
            a["titulo"],
            a["revista"],
            a["doi"],
            a["link"],
            a["base"],
            a["status"],
        ])

    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    thin = Side(border_style="thin", color="D9E2F3")

    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(top=thin, bottom=thin, left=thin, right=thin)

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            cell.border = Border(top=thin, bottom=thin, left=thin, right=thin)

    for row in range(2, ws.max_row + 1):
        doi = ws[f"F{row}"].value
        url = ws[f"G{row}"].value

        if doi:
            ws[f"F{row}"].hyperlink = f"https://doi.org/{doi}"
            ws[f"F{row}"].style = "Hyperlink"

        if url:
            ws[f"G{row}"].hyperlink = str(url)
            ws[f"G{row}"].style = "Hyperlink"

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions

    widths = {
        "A": 8,
        "B": 18,
        "C": 10,
        "D": 55,
        "E": 32,
        "F": 30,
        "G": 42,
        "H": 14,
        "I": 14,
    }

    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    wb.save(caminho)
    return str(caminho)


def exportar_word_tabela_artigo(artigos, pasta_saida="outputs"):
    Path(pasta_saida).mkdir(parents=True, exist_ok=True)
    caminho = Path(pasta_saida) / "ATHENA_PRISMA_Tabela_Artigo.docx"

    doc = Document()

    titulo = doc.add_heading("ATHENA PRISMA REVIEW ROBOT", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = doc.add_heading("Tabela de Estudos Incluídos", level=2)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(f"Estudos incluídos (n = {len(artigos)}).")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = "Table Grid"

    headers = ["Autor", "Ano", "Título", "Revista"]
    for i, h in enumerate(headers):
        tabela.rows[0].cells[i].text = h

    for artigo in artigos:
        a = normalizar_artigo(artigo)
        row = tabela.add_row().cells
        row[0].text = str(a["autor"])
        row[1].text = str(a["ano"])
        row[2].text = str(a["titulo"])
        row[3].text = str(a["revista"])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

    doc.save(caminho)
    return str(caminho)


def exportar_master_prisma(artigos_totais, artigos_incluidos, svg_fluxo=None, jpg_fluxo=None, pdf_fluxo=None):
    arquivos = {}

    arquivos["excel_master"] = exportar_excel_master(artigos_incluidos)
    arquivos["word_tabela_artigo"] = exportar_word_tabela_artigo(artigos_incluidos)

    if svg_fluxo:
        arquivos["fluxograma_svg"] = svg_fluxo
    if pdf_fluxo:
        arquivos["fluxograma_pdf"] = pdf_fluxo

    return arquivos
